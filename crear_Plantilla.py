import streamlit as st

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_plantilla():
    st.title("Crea tu Plantilla Personalizada")
    st.write("Personaliza los elementos de tu informe y descarga tu plantilla.")

    fuente = st.selectbox("Elige la fuente", ["Arial", "Times New Roman", "Calibri"])
    tamano_fuente = st.slider("Tamaño de la fuente", 8, 32, 12)
    color_texto = st.color_picker("Color del texto", "#000000")

    encabezado = st.text_area("Encabezado", "Encabezado de ejemplo")
    pie_pagina = st.text_area("Pie de Página", "Pie de página de ejemplo")
    cuerpo_texto = st.text_area("Texto principal", "Aquí escribe el contenido de tu documento.")

    if st.button("Generar y Descargar Plantilla"):
        doc = Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        encabezado_paragraph = doc.add_paragraph()
        encabezado_run = encabezado_paragraph.add_run(encabezado)
        encabezado_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        encabezado_run.font.name = fuente
        encabezado_run.font.size = Pt(tamano_fuente)
        encabezado_run.font.color.rgb = RGBColor.from_string(color_texto[1:])

        doc.add_paragraph("\n")
        cuerpo_paragraph = doc.add_paragraph()
        cuerpo_run = cuerpo_paragraph.add_run(cuerpo_texto)
        cuerpo_run.font.name = fuente
        cuerpo_run.font.size = Pt(tamano_fuente)
        cuerpo_run.font.color.rgb = RGBColor.from_string(color_texto[1:])

        footer = sections[0].footer.paragraphs[0]
        footer_paragraph = footer.add_run(pie_pagina)
        footer_paragraph.font.name = fuente
        footer_paragraph.font.size = Pt(tamano_fuente)
        footer_paragraph.font.color.rgb = RGBColor.from_string(color_texto[1:])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="Descargar Plantilla",
            data=buffer,
            file_name="plantilla_personalizada.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        buffer.close()

def generar_pdf(contenido):
    html_content = f"<h1>Informe de Ejemplo</h1><p>{contenido}</p>"
    pdf_buffer = BytesIO()
    with open("temp.pdf", "wb") as f:
        f.write(pdfkit.from_string(html_content, False))
    pdf_buffer.seek(0)
    return pdf_buffer
