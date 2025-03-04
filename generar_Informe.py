import streamlit as st
from docx import Document
from io import BytesIO
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_docx():
    st.title("Generador de Informe en Word")
    st.write("Crea y descarga un informe personalizado en formato Word.")

    # Entrada de encabezado, texto principal y pie de página
    encabezado = st.text_area("Encabezado del Informe", "Encabezado de ejemplo")
    texto_principal = st.text_area("Texto Principal del Informe", "Aquí escribe el contenido del informe.")
    pie_pagina = st.text_area("Pie de Página", "Pie de página de ejemplo")

    # Opciones de estilo de texto
    fuente = st.selectbox("Elige la fuente", ["Arial", "Times New Roman", "Calibri"])
    tamano_fuente = st.slider("Tamaño de la fuente", 8, 32, 12)
    color_texto = st.color_picker("Color del texto", "#000000")

    # Botón para generar el documento Word
    if st.button("Generar y Descargar Informe"):
        # Crear documento Word
        doc = Document()

        # Encabezado
        encabezado_paragraph = doc.add_paragraph()
        encabezado_run = encabezado_paragraph.add_run(encabezado)
        encabezado_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        encabezado_run.font.name = fuente
        encabezado_run.font.size = Pt(tamano_fuente)
        encabezado_run.font.color.rgb = RGBColor.from_string(color_texto[1:])
        
        # Añadir espacio entre el encabezado y el contenido principal
        doc.add_paragraph("\n")

        # Texto principal
        texto_paragraph = doc.add_paragraph()
        texto_run = texto_paragraph.add_run(texto_principal)
        texto_run.font.name = fuente
        texto_run.font.size = Pt(tamano_fuente)
        texto_run.font.color.rgb = RGBColor.from_string(color_texto[1:])

        # Pie de página
        sections = doc.sections
        footer = sections[0].footer.paragraphs[0]
        pie_paragraph = footer.add_run(pie_pagina)
        pie_paragraph.font.name = fuente
        pie_paragraph.font.size = Pt(tamano_fuente)
        pie_paragraph.font.color.rgb = RGBColor.from_string(color_texto[1:])

        # Guardar el archivo en BytesIO para descarga
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Botón de descarga
        st.download_button(
            label="Descargar Informe en Word",
            data=buffer,
            file_name="informe_personalizado.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
